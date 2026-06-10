"""Lectura de documentos de distintos formatos."""

from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class Document:
    doc_id: str
    source_file: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def generate_doc_id(path: Path) -> str:
    """Genera un ID determinista basado en la ruta del archivo."""
    return hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:12]


def _base_metadata(path: Path) -> dict[str, Any]:
    return {
        "source_file": str(path),
        "file_type": path.suffix.lstrip(".").lower(),
        "file_name": path.name,
    }


def load_text_file(path: Path) -> Document:
    """Carga un archivo de texto plano o Markdown."""
    text = path.read_text(encoding="utf-8")
    return Document(
        doc_id=generate_doc_id(path),
        source_file=str(path),
        text=text,
        metadata=_base_metadata(path),
    )


def load_pdf_file(path: Path) -> Document:
    """Extrae texto de un PDF usando pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("Instala pypdf: pip install pypdf") from e

    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        # Marca de página para poder rastrear el origen
        parts.append(f"\n[PAGE:{i}]\n{page_text}")
    text = "\n".join(parts).strip()

    return Document(
        doc_id=generate_doc_id(path),
        source_file=str(path),
        text=text,
        metadata={**_base_metadata(path), "num_pages": len(reader.pages)},
    )


def load_docx_file(path: Path) -> Document:
    """Extrae texto de un DOCX usando python-docx."""
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as e:
        raise ImportError("Instala python-docx: pip install python-docx") from e

    docx = DocxDocument(str(path))
    parts = [p.text for p in docx.paragraphs if p.text]
    text = "\n".join(parts)

    return Document(
        doc_id=generate_doc_id(path),
        source_file=str(path),
        text=text,
        metadata=_base_metadata(path),
    )


def load_document(path: str | Path) -> Document:
    """Selecciona el loader adecuado según la extensión."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"No existe el archivo: {p}")

    suffix = p.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return load_text_file(p)
    if suffix == ".pdf":
        return load_pdf_file(p)
    if suffix in {".docx"}:
        return load_docx_file(p)

    raise ValueError(f"Extensión no soportada: {suffix}")
